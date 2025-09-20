import re
import tempfile
import os
from typing import Optional

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

class ResumeParser:
    """Extract and parse text from resume files (PDF and DOCX)"""
    
    def __init__(self):
        self.supported_formats = []
        
        if PYMUPDF_AVAILABLE or PDFPLUMBER_AVAILABLE:
            self.supported_formats.append('pdf')
        
        if PYTHON_DOCX_AVAILABLE or DOCX2TXT_AVAILABLE:
            self.supported_formats.append('docx')
    
    def extract_text(self, uploaded_file) -> str:
        """Extract text from uploaded file"""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text(uploaded_file)
            elif file_extension == 'docx':
                return self._extract_docx_text(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        except Exception as e:
            raise Exception(f"Failed to extract text from {uploaded_file.name}: {str(e)}")
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        
        # Save uploaded file to temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            text = ""
            
            # Try PyMuPDF first
            if PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(tmp_path)
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    
                    if text.strip():
                        return self._clean_text(text)
                except Exception:
                    pass
            
            # Fall back to pdfplumber
            if PDFPLUMBER_AVAILABLE and not text.strip():
                try:
                    import pdfplumber
                    with pdfplumber.open(tmp_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except Exception:
                    pass
            
            if not text.strip():
                raise Exception("Could not extract text from PDF")
            
            return self._clean_text(text)
        
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        
        # Save uploaded file to temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            text = ""
            
            # Try python-docx first
            if PYTHON_DOCX_AVAILABLE:
                try:
                    doc = Document(tmp_path)
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                        text += "\n"
                    
                    if text.strip():
                        return self._clean_text(text)
                except Exception:
                    pass
            
            # Fall back to docx2txt
            if DOCX2TXT_AVAILABLE and not text.strip():
                try:
                    text = docx2txt.process(tmp_path)
                except Exception:
                    pass
            
            if not text.strip():
                raise Exception("Could not extract text from DOCX")
            
            return self._clean_text(text)
        
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        if not text:
            return ""
        
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text.strip())
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace again
        
        # Remove header/footer patterns (common resume artifacts)
        header_footer_patterns = [
            r'Page \d+ of \d+',
            r'Resume - .+',
            r'CV - .+',
            r'Confidential',
            r'References available upon request'
        ]
        
        for pattern in header_footer_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def extract_sections(self, text: str) -> dict:
        """Extract different sections from resume text"""
        
        sections = {
            'personal_info': '',
            'education': '',
            'experience': '',
            'skills': '',
            'projects': '',
            'certifications': '',
            'other': ''
        }
        
        # Define section headers patterns
        section_patterns = {
            'education': r'(education|academic|qualification|degree)',
            'experience': r'(experience|employment|work|professional|career)',
            'skills': r'(skills|technical|competencies|expertise)',
            'projects': r'(projects|portfolio|work samples)',
            'certifications': r'(certification|certificates|training|course)'
        }
        
        # Split text into lines
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            line_lower = line.lower()
            section_found = False
            
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_lower) and len(line) < 50:  # Likely a header
                    current_section = section_name
                    section_found = True
                    break
            
            if not section_found:
                sections[current_section] += line + '\n'
        
        # Clean up sections
        for key, value in sections.items():
            sections[key] = value.strip()
        
        return sections
    
    def extract_contact_info(self, text: str) -> dict:
        """Extract contact information from resume text"""
        
        contact_info: Dict[str, Optional[str]] = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\d{10})'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = github_match.group()
        
        return contact_info
